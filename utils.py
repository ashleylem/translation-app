import io
import os
import openai
import docx
import cv2
from cv2 import imdecode
from itertools import groupby
import subprocess
from io import BytesIO
from flask import Flask, request, send_file, send_from_directory, render_template,  url_for, flash, redirect, jsonify, session
from werkzeug.utils import secure_filename
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
import pytesseract
from PIL import Image
from datetime import datetime
from flask_jwt_extended import JWTManager, create_access_token, jwt_required, get_jwt_identity, current_user, set_access_cookies, get_jwt_identity
import calendar
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
from models import User, db
import uuid
import numpy as np



def image_to_docx(file_stream, target_language, source_language, translator_name):
    source_text = ocr_image(file_stream)
    if not source_text.strip():
        raise ValueError("No text detected in the image file")
    
    translated_text = translate_text(source_text, source_language, target_language)

    doc = docx.Document()

    # ...

    # Add the translated text to the Word document
    for paragraph_text in translated_text.split('\n'):
        paragraph = doc.add_paragraph(paragraph_text)
        for run in paragraph.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
    currentDay = datetime.now().day
    currentMonth = datetime.now().month
    currentYear = datetime.now().year
    month_name = calendar.month_name[currentMonth]

    standard_text = f"""CERTIFICATION OF TRANSLATOR'S COMPETENCE
    
I, {translator_name}, hereby certify that the above is an accurate translation of the original Birth Certificate in Spanish, and that I am competent in both English and Spanish to render such translation.

_______________________________________
Translator’s Signature

Immigration Law Office of
Theodore Texidor, Esq.
8095 NW 12th Street Ste. 105
Doral, Florida 33126
State of Florida, County of Miami Dade

Affiant personally known to me    
Sworn before me, this {currentDay}th day of {month_name} {currentYear}

_____________________________________________
Notary Public"""

    for index, line in enumerate(standard_text.split('\n')):
        paragraph = doc.add_paragraph()

        # Apply formatting to the run (text)
        run = paragraph.add_run(line)
        run.font.name = 'Cambria'
        run.font.size = Pt(12)

        # Apply formatting to the paragraph
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(1)
        paragraph_format.space_after = Pt(1)

        if index == 0:  # First line
            run.font.bold = True
            run.font.size = Pt(14)
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif index in [4, 5, 6]:
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif index in list(range(7, 12)):
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run.font.size = Pt(9)
        else:
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    output_stream = io.BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)
    return output_stream

def preprocess_image(file):
    # Convert the file object to a NumPy array
    file_bytes = np.asarray(bytearray(file.read()), dtype=np.uint8)

    # Reset the file object's position to the beginning
    file.seek(0)

    # Decode the file bytes to an OpenCV image
    img = cv2.imdecode(file_bytes, cv2.IMREAD_COLOR)

    img = cv2.resize(img, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC)

    # Convert to grayscale
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

    # Apply a median blur to reduce noise
    dilation = cv2.dilate(gray, np.ones((1,1)))
    # opened = cv2.morphologyEx(gray, cv2.MORPH_OPEN, np.ones((1, 1), np.uint8))
    # eroding = cv2.morphologyEx(gray, cv2.MORPH_ERODE, np.ones((1, 1)))
    closing = cv2.morphologyEx(
        dilation, cv2.MORPH_CLOSE,  np.ones((1, 1), np.uint8))
   


    # Apply adaptive thresholding
    _, new_img = cv2.threshold(
        closing, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

    # Convert the OpenCV image (NumPy array) to a PIL Image
    pil_img = Image.fromarray(new_img)

    # Specify the complete file path with a unique name and desired extension
    file_path = os.path.join('temp', 'preview.jpg')

    # Save the preprocessed image as a JPEG using PIL
    pil_img.save(file_path)

    return file_path


def ocr_image(file):

    # Perform OCR
    data = pytesseract.image_to_string(
        new_img, config="--psm 3")
    return data


def split_text(text, max_tokens):
    words = text.split()
    chunks = []
    current_chunk = []

    for word in words:
        current_chunk.append(word)
        if sum(len(w) for w in current_chunk) >= max_tokens:
            chunks.append(' '.join(current_chunk))
            current_chunk = []

    if current_chunk:
        chunks.append(' '.join(current_chunk))

    return chunks


def translate_text(source_text, source_lang, target_lang, max_tokens_per_chunk=1500):
    print(len(source_text))
    if len(source_text) > max_tokens_per_chunk:
        text_chunks = split_text(source_text, max_tokens_per_chunk)
        translated_chunks = []
        for chunk in text_chunks:
            print(chunk)
            response = openai.Completion.create(
                engine="text-davinci-003",
                prompt=f"Translate the following text from {source_lang} to {target_lang} in the exact format without summarizing: {chunk}",
                max_tokens=max_tokens_per_chunk*2,
                n=1,
                stop=None,
                temperature=0.5,
            )
            translated_chunk = response.choices[0].text.strip()
            translated_chunks.append(translated_chunk)
            translated_text = ' '.join(translated_chunks)
            return translated_text
    else:
        print(source_text)
        response = openai.Completion.create(
                engine="text-davinci-003",
                prompt=f"Translate the following text from {source_lang} to {target_lang} as close to the original formatting: {source_text}",
                max_tokens=max_tokens_per_chunk * 2,
                n=1,
                stop=None,
                temperature=0.5,
            )
        print(response)
        return response.choices[0].text.strip()



